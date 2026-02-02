"""
文件处理模块
提供文件读取、文本提取、文档生成功能
"""

import io
import json
import zipfile
from datetime import datetime

try:
    from docx import Document
    import pdfplumber
except ImportError:
    raise ImportError(
        "缺少依赖库，请先安装：pip install python-docx pdfplumber"
    )


class FileConfig:
    """文件处理配置"""
    MAX_FILE_SIZE_MB = 50


def load_docx_text(file_obj) -> str:
    """从 DOCX 文件提取文本（包含表格）"""
    try:
        document = Document(file_obj)
        text_parts = []

        # 提取段落文本
        for paragraph in document.paragraphs:
            text_parts.append(paragraph.text)

        # 提取表格文本
        for table in document.tables:
            for row in table.rows:
                row_text = " | ".join([cell.text.strip() for cell in row.cells])
                text_parts.append(row_text)

        # 提取页眉页脚
        for section in document.sections:
            # 页眉
            for paragraph in section.header.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            # 页脚
            for paragraph in section.footer.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

        return "\n".join(text_parts)
    except Exception as e:
        raise ValueError(f"DOCX 文件解析失败：{str(e)}") from e


def load_pdf_text(file_obj) -> str:
    """从 PDF 文件提取文本"""
    try:
        pages = []
        with pdfplumber.open(file_obj) as pdf:
            for page in pdf.pages:
                text = page.extract_text() or ""
                pages.append(text)
        return "\n\n".join(pages)
    except Exception as e:
        raise ValueError(f"PDF 文件解析失败：{str(e)}") from e


def extract_file_text(file_obj) -> str:
    """从各种文件格式提取文本"""
    name = file_obj.name.lower()

    # 检查文件大小
    file_obj.seek(0, io.SEEK_END)
    size_mb = file_obj.tell() / (1024 * 1024)
    file_obj.seek(0)

    if size_mb > FileConfig.MAX_FILE_SIZE_MB:
        raise ValueError(
            f"文件过大（{size_mb:.1f}MB），最大支持 {FileConfig.MAX_FILE_SIZE_MB}MB"
        )

    try:
        if name.endswith(".txt"):
            return file_obj.read().decode("utf-8", errors="ignore")
        elif name.endswith(".docx"):
            return load_docx_text(file_obj)
        elif name.endswith(".pdf"):
            return load_pdf_text(file_obj)
        else:
            raise ValueError("不支持的文件格式，请使用 txt/docx/pdf")
    except Exception as e:
        raise ValueError(f"文件读取失败：{str(e)}") from e


def build_docx_bytes(text: str) -> bytes:
    """构建 DOCX 文件的字节数据（保留空行）"""
    document = Document()
    lines = text.splitlines()

    for line in lines:
        # 保留所有行，包括空行
        if line.strip():  # 有内容的行
            document.add_paragraph(line)
        else:  # 空行
            document.add_paragraph("")  # 添加空段落

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.read()


def build_zip_bundle(
    masked_docx: bytes,
    encrypted_dict: dict,
    stamp: str,
    include_stats: bool = False
) -> bytes:
    """构建打包文件（脱敏文档 + 加密还原文件）"""
    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, "w", compression=zipfile.ZIP_DEFLATED) as bundle:
        bundle.writestr(f"masked_{stamp}.docx", masked_docx)
        bundle.writestr(
            f"restore_{stamp}.json",
            json.dumps(encrypted_dict, ensure_ascii=False, indent=2).encode("utf-8")
        )
    buffer.seek(0)
    return buffer.read()
