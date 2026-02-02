"""
增强的文档处理模块
支持保持原文档的格式、标题、字号等
"""

from dataclasses import dataclass
from typing import List, Tuple
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


@dataclass
class ParagraphData:
    """段落数据类，包含文本和格式信息"""
    text: str
    style_name: str = None
    heading_level: int = None
    is_bold: bool = False
    is_italic: bool = False
    font_size: float = None
    alignment: str = None
    is_title: bool = False
    is_table: bool = False
    table_data: List[List[str]] = None


def extract_docx_with_format(file_obj) -> List[ParagraphData]:
    """从 DOCX 文件提取文本和格式信息"""
    try:
        document = Document(file_obj)
        paragraphs_data = []

        # 提取段落（包含格式）
        for paragraph in document.paragraphs:
            # 获取文本
            text = paragraph.text

            # 获取样式信息
            style = paragraph.style
            style_name = style.name if style else None

            # 判断是否为标题
            heading_level = None
            is_title = False
            if style_name:
                if "Heading" in style_name:
                    try:
                        heading_level = int(style_name.split()[-1])
                    except (ValueError, IndexError):
                        heading_level = None
                elif "Title" in style_name:
                    is_title = True

            # 获取字体属性
            runs = paragraph.runs
            is_bold = False
            is_italic = False
            font_size = None

            if runs:
                first_run = runs[0]
                if first_run.font.bold:
                    is_bold = True
                if first_run.font.italic:
                    is_italic = True
                if first_run.font.size:
                    font_size = first_run.font.size.pt

            # 获取对齐方式
            alignment = None
            if paragraph.alignment:
                alignment_map = {
                    0: "left",
                    1: "center",
                    2: "right",
                    3: "justify"
                }
                alignment = alignment_map.get(paragraph.alignment, "left")

            paragraphs_data.append(ParagraphData(
                text=text,
                style_name=style_name,
                heading_level=heading_level,
                is_bold=is_bold,
                is_italic=is_italic,
                font_size=font_size,
                alignment=alignment,
                is_title=is_title
            ))

        # 提取表格
        for table in document.tables:
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text.strip())
                table_data.append(row_data)

            # 将表格作为一个特殊段落
            if table_data:
                table_text = " | ".join([" | ".join(row) for row in table_data])
                paragraphs_data.append(ParagraphData(
                    text=table_text,
                    is_table=True,
                    table_data=table_data
                ))

        # 提取页眉页脚
        for section in document.sections:
            # 页眉
            for paragraph in section.header.paragraphs:
                if paragraph.text.strip():
                    paragraphs_data.append(ParagraphData(
                        text=paragraph.text,
                        style_name="Header"
                    ))

            # 页脚
            for paragraph in section.footer.paragraphs:
                if paragraph.text.strip():
                    paragraphs_data.append(ParagraphData(
                        text=paragraph.text,
                        style_name="Footer"
                    ))

        return paragraphs_data

    except Exception as e:
        raise ValueError(f"DOCX 文件解析失败：{str(e)}") from e


def build_docx_with_format(paragraphs_data: List[ParagraphData]) -> bytes:
    """根据段落数据构建 DOCX 文件（保持格式）"""
    import io
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

    document = Document()

    # 设置默认字体
    document.styles['Normal'].font.name = '宋体'
    document.styles['Normal']._element.rPr.rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', '宋体')

    for para_data in paragraphs_data:
        if not para_data.text.strip():
            # 空段落
            document.add_paragraph()
            continue

        if para_data.is_table and para_data.table_data:
            # 创建表格
            table = document.add_table(rows=len(para_data.table_data), cols=len(para_data.table_data[0]))
            table.style = 'Light Grid Accent 1'

            for i, row_data in enumerate(para_data.table_data):
                for j, cell_text in enumerate(row_data):
                    cell = table.rows[i].cells[j]
                    cell.text = cell_text

            # 添加表格后的空行
            document.add_paragraph()
            continue

        # 创建段落
        para = document.add_paragraph()

        # 设置文本
        para.text = para_data.text

        # 设置样式
        if para_data.heading_level:
            # 标题样式
            heading = document.styles[f'Heading {para_data.heading_level}']
            para.style = heading

        elif para_data.is_title:
            # 标题样式
            title_style = document.styles['Title']
            para.style = title_style

        # 设置格式
        run = para.runs[0] if para.runs else para.add_run(para_data.text)

        if para_data.is_bold:
            run.font.bold = True

        if para_data.is_italic:
            run.font.italic = True

        if para_data.font_size:
            run.font.size = Pt(para_data.font_size)

        # 设置对齐
        if para_data.alignment:
            alignment_map = {
                "left": WD_PARAGRAPH_ALIGNMENT.LEFT,
                "center": WD_PARAGRAPH_ALIGNMENT.CENTER,
                "right": WD_PARAGRAPH_ALIGNMENT.RIGHT,
                "justify": WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            }
            para.alignment = alignment_map.get(para_data.alignment, WD_PARAGRAPH_ALIGNMENT.LEFT)

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.read()


def mask_paragraphs_data(
    paragraphs_data: List[ParagraphData],
    masked_text: str,
    original_text: str
) -> List[ParagraphData]:
    """
    对段落数据进行脱敏，同时保持格式

    Args:
        paragraphs_data: 原始段落数据列表
        masked_text: 脱敏后的完整文本
        original_text: 原始完整文本

    Returns:
        脱敏后的段落数据列表
    """
    # 将原始文本按段落分割
    original_paragraphs = original_text.split('\n')
    masked_paragraphs = masked_text.split('\n')

    # 如果段落数量不匹配，按行重建
    if len(original_paragraphs) != len(masked_paragraphs):
        # 简单按行重建
        result = []
        for para_data in paragraphs_data:
            # 在脱敏文本中查找对应段落
            new_para = ParagraphData(
                text=mask_text_in_context(para_data.text, masked_text),
                style_name=para_data.style_name,
                heading_level=para_data.heading_level,
                is_bold=para_data.is_bold,
                is_italic=para_data.is_italic,
                font_size=para_data.font_size,
                alignment=para_data.alignment,
                is_title=para_data.is_title,
                is_table=para_data.is_table,
                table_data=para_data.table_data
            )
            result.append(new_para)
        return result

    # 段落数量匹配，直接替换文本
    result = []
    for i, para_data in enumerate(paragraphs_data):
        if i < len(masked_paragraphs):
            new_para = ParagraphData(
                text=masked_paragraphs[i] if para_data.text else "",
                style_name=para_data.style_name,
                heading_level=para_data.heading_level,
                is_bold=para_data.is_bold,
                is_italic=para_data.is_italic,
                font_size=para_data.font_size,
                alignment=para_data.alignment,
                is_title=para_data.is_title,
                is_table=para_data.is_table,
                table_data=para_data.table_data
            )
            result.append(new_para)
        else:
            result.append(para_data)

    return result


def mask_text_in_context(original_text: str, masked_full_text: str) -> str:
    """
    在完整脱敏文本中找到对应段落的脱敏版本

    Args:
        original_text: 原始段落文本
        masked_full_text: 完整的脱敏文本

    Returns:
        该段落的脱敏版本
    """
    # 简单实现：在脱敏文本中查找最相似的段落
    masked_lines = masked_full_text.split('\n')

    # 找到长度最接近且未完全脱敏的行
    best_match = ""
    min_diff = float('inf')

    original_stripped = original_text.strip()
    original_len = len(original_stripped)

    for line in masked_lines:
        line_stripped = line.strip()

        # 跳过空行
        if not line_stripped:
            if not original_stripped:
                return ""
            continue

        # 跳过完全脱敏的行（全是*号）
        if '*' in line_stripped and len([c for c in line_stripped if c != '*']) < 3:
            continue

        # 找到长度最接近的
        diff = abs(len(line_stripped) - original_len)
        if diff < min_diff:
            min_diff = diff
            best_match = line_stripped

    return best_match if best_match else original_text
