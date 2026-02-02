"""
高级 OOXML 文档处理器
在 Run 级别处理 Word 文档，保持样式和布局稳定性
"""

import re
from typing import List, Dict, Tuple, Optional
from docx import Document
from docx.text.paragraph import Run
from dataclasses import dataclass
from lxml import etree


@dataclass
class TextReplacement:
    """文本替换指令"""
    start: int
    end: int
    original: str
    replacement: str
    # Equal-length replacement to preserve layout
    use_equal_length: bool = True


class OOXMLProcessor:
    """
    OOXML 文档处理器
    在 Run 级别进行文本替换，保持样式和布局
    """

    def __init__(self, document: Document):
        self.document = document
        self.ns = {
            'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        }

    def process_document(self, replacements: List[Tuple[str, str, bool]]) -> Document:
        """
        处理整个文档，应用所有替换规则

        Args:
            replacements: 替换规则列表 [(pattern, replacement, use_equal_length), ...]

        Returns:
            处理后的文档
        """
        # 处理正文段落
        for paragraph in self.document.paragraphs:
            self._process_paragraph(paragraph, replacements)

        # 处理表格（包括嵌套表格）
        for table in self.document.tables:
            self._process_table(table, replacements)

        # 处理页眉
        for section in self.document.sections:
            for paragraph in section.header.paragraphs:
                self._process_paragraph(paragraph, replacements)
            # 页眉中的表格
            for table in section.header.tables:
                self._process_table(table, replacements)

        # 处理页脚
        for section in self.document.sections:
            for paragraph in section.footer.paragraphs:
                self._process_paragraph(paragraph, replacements)
            # 页脚中的表格
            for table in section.footer.tables:
                self._process_table(table, replacements)

        return self.document

    def _process_paragraph(self, paragraph, replacements: List[Tuple[str, str, bool]]):
        """
        处理单个段落，在 Run 级别进行文本替换

        关键逻辑：
        1. 合并段落中所有 Run 的文本
        2. 在合并后的文本中查找需要替换的内容
        3. 将替换结果写回原始的 Run 结构中，保持样式
        """
        # 收集所有 Run 的文本和样式信息
        runs_data = []
        full_text = ""

        for run in paragraph.runs:
            run_text = run.text
            runs_data.append({
                'run': run,
                'text': run_text,
                'length': len(run_text)
            })
            full_text += run_text

        if not full_text:
            return

        # 在完整文本中应用所有替换规则
        masked_text = full_text
        replacement_map = []  # 记录所有替换的位置

        for pattern, replacement_char, use_equal_length in replacements:
            if use_equal_length:
                # 等长替换策略
                matches = re.finditer(pattern, full_text)
                for match in reversed(list(matches)):  # 从后往前替换，避免位置偏移
                    start, end = match.span()
                    original = match.group()

                    # 创建等长占位符
                    placeholder = self._create_equal_length_placeholder(
                        original, replacement_char
                    )

                    replacement_map.append(TextReplacement(
                        start=start,
                        end=end,
                        original=original,
                        replacement=placeholder,
                        use_equal_length=True
                    ))
            else:
                # 直接替换
                masked_text = re.sub(pattern, lambda m: replacement_char * len(m.group()), masked_text)

        # 如果有等长替换，需要精确应用
        if replacement_map:
            # 按位置排序
            replacement_map.sort(key=lambda x: x.start)

            # 构建最终文本
            masked_text = self._apply_replacements(full_text, replacement_map)

        # 将处理后的文本写回 Run 结构
        self._write_text_to_runs(paragraph, runs_data, full_text, masked_text)

    def _process_table(self, table, replacements: List[Tuple[str, str, bool]]):
        """
        处理表格（包括嵌套表格）

        Args:
            table: Word 表格对象
            replacements: 替换规则列表
        """
        for row in table.rows:
            for cell in row.cells:
                # 处理单元格中的段落
                for paragraph in cell.paragraphs:
                    self._process_paragraph(paragraph, replacements)

                # 递归处理嵌套表格
                for nested_table in cell.tables:
                    self._process_table(nested_table, replacements)

    def _create_equal_length_placeholder(self, original: str, mask_char: str = "█") -> str:
        """
        创建等长占位符，保持布局稳定

        使用全角字符或特殊 Unicode 字符来保持视觉宽度

        Args:
            original: 原始文本
            mask_char: 掩码字符

        Returns:
            与原始文本等长的占位符
        """
        # 分析原始文本的字符类型
        full_width_chars = 0
        half_width_chars = 0

        for char in original:
            # 判断是否为全角字符（中文、全角标点等）
            if '\u4e00' <= char <= '\u9fff' or char in '，。、；：""''（）【】《》':
                full_width_chars += 1
            else:
                half_width_chars += 1

        # 创建等长占位符
        # 使用全角方块字符（U+2588）作为全角占位符
        # 使用普通星号作为半角占位符
        placeholder = ""
        placeholder += "█" * full_width_chars  # 全角占位符
        placeholder += "*" * half_width_chars  # 半角占位符

        return placeholder

    def _apply_replacements(self, text: str, replacements: List[TextReplacement]) -> str:
        """
        应用所有替换规则到文本

        Args:
            text: 原始文本
            replacements: 替换指令列表

        Returns:
            替换后的文本
        """
        # 转换为列表以便修改
        result = list(text)

        # 从后往前应用替换（避免位置偏移）
        for repl in reversed(replacements):
            # 替换文本
            for i in range(repl.start, repl.end):
                if i < len(result):
                    # 计算在 replacement 中的位置
                    pos_in_replacement = i - repl.start
                    if pos_in_replacement < len(repl.replacement):
                        result[i] = repl.replacement[pos_in_replacement]

        return ''.join(result)

    def _write_text_to_runs(self, paragraph, runs_data: List[Dict], original_text: str, masked_text: str):
        """
        将处理后的文本写回 Run 结构，保持样式

        这是关键逻辑：需要智能地分配文本到各个 Run，保持样式不变

        Args:
            paragraph: 段落对象
            runs_data: Run 数据列表（包含原始文本和 Run 对象）
            original_text: 原始完整文本
            masked_text: 处理后的完整文本
        """
        # 如果长度差异太大（等长替换应该保持长度一致），需要特殊处理
        if len(original_text) != len(masked_text):
            # 长度不一致，无法保持精确的 Run 结构
            # 使用备用策略：清空所有 Run，在第一个 Run 中放入完整文本
            if paragraph.runs:
                # 保存第一个 Run 的样式
                first_run = paragraph.runs[0]
                # 清空所有 Run
                for run in paragraph.runs:
                    run.text = ""
                # 在第一个 Run 中设置完整文本
                first_run.text = masked_text
            return

        # 长度一致，精确分配文本到各个 Run
        current_pos = 0

        for i, run_data in enumerate(runs_data):
            run = run_data['run']
            original_length = run_data['length']

            if original_length == 0:
                run.text = ""
                continue

            # 计算这个 Run 应该对应的文本范围
            end_pos = current_pos + original_length

            if end_pos > len(masked_text):
                # 边界情况
                run.text = masked_text[current_pos:]
            else:
                run.text = masked_text[current_pos:end_pos]

            current_pos = end_pos


def apply_ooxml_masking(
    file_obj,
    keywords: List[str],
    mask_patterns: Dict[str, str],
    preserve_suffix: bool = True
) -> bytes:
    """
    应用 OOXML 级别的脱敏处理

    Args:
        file_obj: 文件对象
        keywords: 关键词列表
        mask_patterns: 脱敏模式字典 {名称: 正则表达式}
        preserve_suffix: 是否保留公司后缀

    Returns:
        处理后的 DOCX 文件字节
    """
    import io

    # 加载文档
    document = Document(file_obj)

    # 创建处理器
    processor = OOXMLProcessor(document)

    # 构建替换规则列表
    replacements = []

    # 添加关键词替换规则（使用等长占位符）
    for keyword in keywords:
        escaped = re.escape(keyword)
        replacements.append((escaped, "*", True))

    # 添加智能识别规则
    # 手机号 - 等长替换
    replacements.append((r"1[3-9]\d{9}", "*", True))

    # 身份证号 - 等长替换
    replacements.append((
        r"[1-9]\d{5}(18|19|20)\d{2}(0[1-9]|1[0-2])(0[1-9]|[12]\d|3[01])\d{3}[\dXx]",
        "*",
        True
    ))

    # 邮箱 - 等长替换
    replacements.append((r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b", "*", True))

    # 银行卡号 - 等长替换
    replacements.append((r"\b\d{16,19}\b", "*", True))

    # IP地址 - 等长替换
    replacements.append((r"\b\d{1,3}\.\d{1,3}\.\d{1,3}\.\d{1,3}\b", "*", True))

    # 统一社会信用代码 - 等长替换
    replacements.append((r"[0-9A-HJ-NPQRTUW-Y]{2}\d{6}[0-9A-HJ-NPQRTUW-Y]{10}", "*", True))

    # 企业名称 - 特殊处理（保留后缀）
    if preserve_suffix:
        # 企业名称需要特殊处理，不在等长替换中处理
        # 这个会在后面单独处理
        pass

    # 详细地址 - 等长替换
    replacements.append((
        r"[\u4e00-\u9fa5]{2,6}(?:省|市|区|县|镇|乡|街道|路|巷|号|栋|单元|楼|层|室|户)[\u4e00-\u9fa5\d\-#号]*",
        "*",
        True
    ))

    # 车牌号 - 等长替换
    replacements.append((
        r"[京津沪渝冀豫云辽黑湘皖鲁新苏浙赣鄂桂甘晋蒙陕吉闽贵粤青藏川宁琼使领][A-Z][A-Z0-9]{5,6}",
        "*",
        True
    ))

    # 金额 - 等长替换（完全脱敏）
    replacements.append((
        r"(?:¥|￥|USD?|\$)\s*(?:\d{1,3}(?:,\d{3})+(?:\.\d{1,2})?|\d+(?:\.\d{1,2})?)\s*(?:万元?|元)?|(?:\d{1,3}(?:,\d{3})+|\d+)(?:\.\d{1,2})?\s*(?:万元?|元)",
        "*",
        True
    ))

    # 应用所有替换
    processed_doc = processor.process_document(replacements)

    # 特殊处理：企业名称脱敏（保留后缀）
    # 这个需要在文档级别单独处理，因为涉及到后缀保留逻辑
    if preserve_suffix:
        _apply_company_masking(processed_doc)

    # 保存到字节流
    buffer = io.BytesIO()
    processed_doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


def _apply_company_masking(document: Document):
    """
    对文档应用企业名称脱敏（保留后缀）

    Args:
        document: Word 文档对象
    """
    # 修正后的企业名称模式：匹配任意长度的企业名称
    # 修改：使用更广泛的字符范围和更大的长度限制
    company_pattern = re.compile(
        r"[\u4e00-\u9fa5a-zA-Z0-9()（）]{2,30}(?:有限公司|股份有限公司|有限责任公司|集团有限公司|公司)"
    )

    # 定义后缀列表（从长到短排序，确保优先匹配长的后缀）
    suffixes = [
        "股份有限公司",
        "有限责任公司",
        "集团有限公司",
        "有限公司",
        "公司",
    ]

    def mask_company_name(match):
        """脱敏企业名称，保留后缀"""
        original = match.group(0)

        # 找到后缀（从长到短尝试）
        for suffix in suffixes:
            if original.endswith(suffix):
                name_part = original[:-len(suffix)]
                suffix_part = suffix
                # 使用等长占位符脱敏名称部分
                # 注意：len() 返回字符数，对于中文字符是准确的
                return "█" * len(name_part) + suffix_part

        # 如果没有匹配到后缀，完全脱敏
        return "█" * len(original)

    # 处理所有文本内容
    for paragraph in document.paragraphs:
        new_text = company_pattern.sub(mask_company_name, paragraph.text)
        if paragraph.runs:
            # 保留第一个 Run 的样式
            first_run = paragraph.runs[0]
            for run in paragraph.runs:
                run.text = ""
            first_run.text = new_text

    # 处理表格
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    new_text = company_pattern.sub(mask_company_name, paragraph.text)
                    if paragraph.runs:
                        first_run = paragraph.runs[0]
                        for run in paragraph.runs:
                            run.text = ""
                        first_run.text = new_text

                # 处理嵌套表格
                for nested_table in cell.tables:
                    for nested_row in nested_table.rows:
                        for nested_cell in nested_row.cells:
                            for paragraph in nested_cell.paragraphs:
                                new_text = company_pattern.sub(mask_company_name, paragraph.text)
                                if paragraph.runs:
                                    first_run = paragraph.runs[0]
                                    for run in paragraph.runs:
                                        run.text = ""
                                    first_run.text = new_text

    # 处理页眉页脚
    for section in document.sections:
        for paragraph in section.header.paragraphs:
            new_text = company_pattern.sub(mask_company_name, paragraph.text)
            if paragraph.runs:
                first_run = paragraph.runs[0]
                for run in paragraph.runs:
                    run.text = ""
                first_run.text = new_text

        for paragraph in section.footer.paragraphs:
            new_text = company_pattern.sub(mask_company_name, paragraph.text)
            if paragraph.runs:
                first_run = paragraph.runs[0]
                for run in paragraph.runs:
                    run.text = ""
                first_run.text = new_text
