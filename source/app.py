"""
自定义脱敏工具 - 定义词脱敏与加密还原
支持多种脱敏模式：全量替换、部分遮蔽、正则匹配、智能识别
所有处理在本地完成，保护数据隐私
"""

import base64
import io
import json
import os
import re
import zipfile
from dataclasses import dataclass, asdict
from datetime import datetime
from enum import Enum
from typing import Callable, Optional

import streamlit as st

# 文件处理导入（延迟导入优化）
try:
    from cryptography.hazmat.primitives.ciphers.aead import AESGCM
    from cryptography.hazmat.primitives.kdf.pbkdf2 import PBKDF2HMAC
    from cryptography.hazmat.primitives import hashes
    from docx import Document
    import pdfplumber
except ImportError as exc:
    st.error("❌ 缺少依赖库，请先安装：pip install -r requirements.txt")
    raise exc


# ============= 常量定义 =============

class Config:
    """应用配置常量"""
    PAGE_TITLE = "定义词脱敏工具"
    PAGE_LAYOUT = "wide"

    # 加密配置
    SALT_LENGTH = 16
    NONCE_LENGTH = 12
    PBKDF2_ITERATIONS = 120000
    KEY_LENGTH = 32

    # UI 配置
    DEFAULT_TEXT_SAMPLE = (
        "客户姓名：张三\n"
        "身份证号：110101199003071234\n"
        "手机号：13800138000\n"
        "邮箱：zhangsan@example.com\n"
        "银行卡号：6222021234567890123\n"
        "住址：北京市朝阳区望京街道 88 号\n"
        "合同金额：¥128,000 元"
    )

    # 文件大小限制（MB）
    MAX_FILE_SIZE_MB = 50


class MaskMode(Enum):
    """脱敏模式枚举"""
    FULL = "full"  # 全量替换：张三 -> ***
    PARTIAL = "partial"  # 部分遮蔽：张三 -> 张*
    REGEX = "regex"  # 正则匹配
    SMART = "smart"  # 智能识别


# ============= 数据模型 =============

@dataclass
class MaskPattern:
    """脱敏模式配置"""
    name: str
    pattern: str
    mode: MaskMode
    preserve_chars: int = 0
    mask_char: str = "*"
    description: str = ""


@dataclass
class EncryptionResult:
    """加密结果数据类"""
    salt: str
    nonce: str
    data: str
    created_at: str
    version: str
    original_length: int
    masked_keywords: list[str]


# ============= 预定义脱敏模式 =============

PREDEFINED_PATTERNS = {
    "手机号": MaskPattern(
        name="手机号",
        pattern=r"1[3-9]\d{9}",
        mode=MaskMode.PARTIAL,
        preserve_chars=3,
        mask_char="*",
        description="中国大陆手机号，保留前3位"
    ),
    "身份证号": MaskPattern(
        name="身份证号",
        pattern=r"[1-9]\d{5}(18|19|20)\d{2}(0[1-9]|1[0-2])(0[1-9]|[12]\d|3[01])\d{3}[\dXx]",
        mode=MaskMode.PARTIAL,
        preserve_chars=6,
        mask_char="*",
        description="18位身份证号，保留前6位"
    ),
    "邮箱": MaskPattern(
        name="邮箱",
        pattern=r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b",
        mode=MaskMode.PARTIAL,
        preserve_chars=2,
        mask_char="*",
        description="电子邮箱地址，保留前2位"
    ),
    "银行卡号": MaskPattern(
        name="银行卡号",
        pattern=r"\b\d{16,19}\b",
        mode=MaskMode.PARTIAL,
        preserve_chars=4,
        mask_char="*",
        description="银行卡号，保留前4位"
    ),
    "IP地址": MaskPattern(
        name="IP地址",
        pattern=r"\b\d{1,3}\.\d{1,3}\.\d{1,3}\.\d{1,3}\b",
        mode=MaskMode.PARTIAL,
        preserve_chars=4,
        mask_char="*",
        description="IPv4地址"
    ),
}


# ============= 核心脱敏功能 =============

def normalize_keywords(raw: str) -> list[str]:
    """解析关键词列表，支持换行、逗号、分号分隔"""
    if not raw:
        return []
    return [item.strip() for item in re.split(r"[\n,;，；]", raw) if item.strip()]


def mask_text_full(text: str, keyword: str, mask_char: str = "*") -> str:
    """全量替换模式"""
    escaped = re.escape(keyword)
    return re.sub(escaped, mask_char * len(keyword), text)


def mask_text_partial(text: str, keyword: str, preserve_chars: int = 1, mask_char: str = "*") -> str:
    """部分遮蔽模式：保留前N位，其余用*替换"""
    escaped = re.escape(keyword)

    def replacement(match):
        original = match.group(0)
        if len(original) <= preserve_chars:
            return original
        return original[:preserve_chars] + mask_char * (len(original) - preserve_chars)

    return re.sub(escaped, replacement, text)


def mask_text_regex(text: str, pattern: str, preserve_chars: int = 0, mask_char: str = "*") -> tuple[str, int]:
    """正则匹配模式"""
    match_count = 0

    def replacement(match):
        nonlocal match_count
        match_count += 1
        original = match.group(0)
        if len(original) <= preserve_chars:
            return original
        return original[:preserve_chars] + mask_char * (len(original) - preserve_chars)

    masked = re.sub(pattern, replacement, text)
    return masked, match_count


def apply_smart_detection(text: str) -> tuple[str, dict[str, int]]:
    """智能识别常见敏感信息并脱敏"""
    stats = {}
    result = text

    for name, pattern in PREDEFINED_PATTERNS.items():
        result, count = mask_text_regex(
            result,
            pattern.pattern,
            pattern.preserve_chars,
            pattern.mask_char
        )
        if count > 0:
            stats[name] = count

    return result, stats


def build_masked_text(
    text: str,
    keywords: list[str],
    mask_mode: MaskMode = MaskMode.FULL,
    preserve_chars: int = 1,
    mask_char: str = "*",
    enable_smart: bool = False
) -> tuple[str, dict]:
    """构建脱敏文本"""
    masked = text
    stats = {"manual_keywords": len(keywords), "smart_detection": {}}

    # 应用手动关键词脱敏
    if keywords:
        if mask_mode == MaskMode.FULL:
            for word in keywords:
                masked = mask_text_full(masked, word, mask_char)
        elif mask_mode == MaskMode.PARTIAL:
            for word in keywords:
                masked = mask_text_partial(masked, word, preserve_chars, mask_char)

    # 应用智能识别
    if enable_smart:
        masked, smart_stats = apply_smart_detection(masked)
        stats["smart_detection"] = smart_stats

    return masked, stats


# ============= 加密解密功能 =============

def derive_key(password: str, salt: bytes) -> bytes:
    """从密码派生加密密钥"""
    kdf = PBKDF2HMAC(
        algorithm=hashes.SHA256(),
        length=Config.KEY_LENGTH,
        salt=salt,
        iterations=Config.PBKDF2_ITERATIONS,
    )
    return kdf.derive(password.encode("utf-8"))


def encrypt_text(text: str, password: str, keywords: list[str]) -> EncryptionResult:
    """加密原文并保存元数据"""
    salt = os.urandom(Config.SALT_LENGTH)
    key = derive_key(password, salt)
    aesgcm = AESGCM(key)
    nonce = os.urandom(Config.NONCE_LENGTH)
    data = aesgcm.encrypt(nonce, text.encode("utf-8"), None)

    return EncryptionResult(
        salt=base64.b64encode(salt).decode("utf-8"),
        nonce=base64.b64encode(nonce).decode("utf-8"),
        data=base64.b64encode(data).decode("utf-8"),
        created_at=datetime.now().isoformat(timespec="seconds"),
        version="2.0",
        original_length=len(text),
        masked_keywords=keywords
    )


def decrypt_text(payload: dict, password: str) -> str:
    """解密还原原文"""
    try:
        salt = base64.b64decode(payload["salt"])
        nonce = base64.b64decode(payload["nonce"])
        data = base64.b64decode(payload["data"])
        key = derive_key(password, salt)
        aesgcm = AESGCM(key)
        plain = aesgcm.decrypt(nonce, data, None)
        return plain.decode("utf-8")
    except (KeyError, ValueError) as e:
        raise ValueError("加密文件格式错误或已损坏") from e


# ============= 文件处理功能 =============

def load_docx_text(file_obj) -> str:
    """从 DOCX 文件提取文本"""
    try:
        document = Document(file_obj)
        return "\n".join([paragraph.text for paragraph in document.paragraphs])
    except Exception as e:
        raise ValueError(f"DOCX 文件解析失败：{str(e)}") from e


def load_pdf_text(file_obj) -> str:
    """从 PDF 文件提取文本"""
    try:
        pages = []
        with pdfplumber.open(file_obj) as pdf:
            for page in pdf.pages:
                text = page.extract_text() or ""
                pages.append(text)
        return "\n\n".join(pages)
    except Exception as e:
        raise ValueError(f"PDF 文件解析失败：{str(e)}") from e


def extract_file_text(file_obj) -> str:
    """从各种文件格式提取文本"""
    name = file_obj.name.lower()

    # 检查文件大小
    file_obj.seek(0, os.SEEK_END)
    size_mb = file_obj.tell() / (1024 * 1024)
    file_obj.seek(0)

    if size_mb > Config.MAX_FILE_SIZE_MB:
        raise ValueError(f"文件过大（{size_mb:.1f}MB），最大支持 {Config.MAX_FILE_SIZE_MB}MB")

    try:
        if name.endswith(".txt"):
            return file_obj.read().decode("utf-8", errors="ignore")
        elif name.endswith(".docx"):
            return load_docx_text(file_obj)
        elif name.endswith(".pdf"):
            return load_pdf_text(file_obj)
        else:
            raise ValueError("不支持的文件格式，请使用 txt/docx/pdf")
    except Exception as e:
        raise ValueError(f"文件读取失败：{str(e)}") from e


def build_docx_bytes(text: str) -> bytes:
    """构建 DOCX 文件的字节数据"""
    document = Document()
    for line in text.splitlines():
        if line.strip():
            document.add_paragraph(line)
    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.read()


def build_zip_bundle(
    masked_docx: bytes,
    encrypted_dict: dict,
    stamp: str,
    include_stats: bool = False
) -> bytes:
    """构建打包文件（脱敏文档 + 加密还原文件）"""
    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, "w", compression=zipfile.ZIP_DEFLATED) as bundle:
        bundle.writestr(f"masked_{stamp}.docx", masked_docx)
        bundle.writestr(
            f"restore_{stamp}.json",
            json.dumps(encrypted_dict, ensure_ascii=False, indent=2).encode("utf-8")
        )
    buffer.seek(0)
    return buffer.read()


# ============= Streamlit UI =============

def init_page_style():
    """初始化页面样式"""
    st.set_page_config(
        page_title=Config.PAGE_TITLE,
        layout=Config.PAGE_LAYOUT,
        page_icon="🔒"
    )

    st.markdown(
        """
        <style>
        @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700&family=JetBrains+Mono:wght@400;500&display=swap');

        :root {
            --bg-primary: #0a0e1a;
            --bg-secondary: #111827;
            --bg-card: rgba(17, 24, 39, 0.7);
            --bg-elevated: rgba(31, 41, 55, 0.5);
            --border-color: rgba(99, 102, 241, 0.2);
            --border-hover: rgba(99, 102, 241, 0.4);
            --accent-primary: #818cf8;
            --accent-secondary: #6366f1;
            --accent-glow: rgba(99, 102, 241, 0.15);
            --text-primary: #f3f4f6;
            --text-secondary: #9ca3af;
            --text-muted: #6b7280;
            --success: #34d399;
            --warning: #fbbf24;
            --error: #f87171;
            --shadow-sm: 0 1px 2px 0 rgba(0, 0, 0, 0.3);
            --shadow-md: 0 4px 6px -1px rgba(0, 0, 0, 0.4);
            --shadow-lg: 0 10px 15px -3px rgba(0, 0, 0, 0.5);
            --shadow-glow: 0 0 20px rgba(99, 102, 241, 0.3);
        }

        * {
            box-sizing: border-box;
        }

        html, body, .stApp {
            font-family: "Inter", -apple-system, BlinkMacSystemFont, sans-serif;
            color: var(--text-primary);
            background: var(--bg-primary);
        }

        body::before {
            content: '';
            position: fixed;
            top: 0;
            left: 0;
            right: 0;
            bottom: 0;
            background:
                radial-gradient(circle at 20% 20%, rgba(99, 102, 241, 0.08) 0%, transparent 50%),
                radial-gradient(circle at 80% 80%, rgba(139, 92, 246, 0.06) 0%, transparent 50%),
                radial-gradient(circle at 40% 60%, rgba(59, 130, 246, 0.05) 0%, transparent 50%);
            pointer-events: none;
            z-index: -1;
        }

        .block-container {
            padding-top: 2.5rem;
            padding-bottom: 2rem;
            max-width: 1400px;
        }

        .app-header {
            text-align: center;
            margin-bottom: 2.5rem;
            animation: fadeInDown 0.6s ease-out;
        }

        .app-header h1 {
            font-size: clamp(28px, 5vw, 42px);
            font-weight: 700;
            background: linear-gradient(135deg, var(--accent-primary) 0%, var(--accent-secondary) 100%);
            -webkit-background-clip: text;
            -webkit-text-fill-color: transparent;
            background-clip: text;
            margin: 0 0 0.5rem 0;
            letter-spacing: -0.02em;
        }

        .app-header p {
            color: var(--text-secondary);
            font-size: 16px;
            margin: 0;
            max-width: 600px;
            margin-left: auto;
            margin-right: auto;
            line-height: 1.6;
        }

        .card {
            background: var(--bg-card);
            border: 1px solid var(--border-color);
            border-radius: 16px;
            padding: 1.5rem;
            backdrop-filter: blur(20px);
            box-shadow: var(--shadow-md), var(--shadow-glow);
            transition: all 0.3s ease;
            animation: fadeInUp 0.6s ease-out backwards;
        }

        .card:hover {
            border-color: var(--border-hover);
            box-shadow: var(--shadow-lg), var(--shadow-glow);
        }

        .card-title {
            font-size: 1.1rem;
            font-weight: 600;
            color: var(--text-primary);
            margin: 0 0 1rem 0;
            display: flex;
            align-items: center;
            gap: 0.5rem;
        }

        .card-title::before {
            content: '';
            width: 4px;
            height: 18px;
            background: linear-gradient(180deg, var(--accent-primary) 0%, var(--accent-secondary) 100%);
            border-radius: 2px;
        }

        .stTextArea > div > div > textarea,
        .stTextInput > div > div > input,
        .stSelectbox > div > div > select {
            background: var(--bg-elevated);
            border: 1px solid var(--border-color);
            border-radius: 10px;
            color: var(--text-primary);
            font-family: "JetBrains Mono", monospace;
            font-size: 14px;
            transition: all 0.2s ease;
        }

        .stTextArea > div > div > textarea:focus,
        .stTextInput > div > div > input:focus,
        .stSelectbox > div > div > select:focus {
            border-color: var(--accent-secondary);
            box-shadow: 0 0 0 3px var(--accent-glow);
            outline: none;
        }

        .stButton > button {
            background: linear-gradient(135deg, var(--accent-secondary) 0%, var(--accent-primary) 100%);
            color: white;
            border: none;
            border-radius: 10px;
            padding: 0.6rem 1.5rem;
            font-weight: 600;
            font-size: 15px;
            transition: all 0.2s ease;
            box-shadow: 0 4px 12px rgba(99, 102, 241, 0.3);
        }

        .stButton > button:hover {
            transform: translateY(-1px);
            box-shadow: 0 6px 16px rgba(99, 102, 241, 0.4);
        }

        .stButton > button:active {
            transform: translateY(0);
        }

        .secondary-btn .stButton > button {
            background: linear-gradient(135deg, var(--success) 0%, #10b981 100%);
            box-shadow: 0 4px 12px rgba(52, 211, 153, 0.3);
        }

        .info-box {
            background: var(--bg-elevated);
            border-left: 3px solid var(--accent-primary);
            border-radius: 8px;
            padding: 1rem;
            margin: 1rem 0;
            font-size: 0.9rem;
            color: var(--text-secondary);
        }

        .success-text {
            color: var(--success);
            font-weight: 500;
        }

        .error-text {
            color: var(--error);
            font-weight: 500;
        }

        .stats-grid {
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(120px, 1fr));
            gap: 0.8rem;
            margin: 1rem 0;
        }

        .stat-item {
            background: var(--bg-elevated);
            border-radius: 10px;
            padding: 0.8rem;
            text-align: center;
            border: 1px solid var(--border-color);
        }

        .stat-value {
            font-size: 1.5rem;
            font-weight: 700;
            color: var(--accent-primary);
        }

        .stat-label {
            font-size: 0.8rem;
            color: var(--text-muted);
            margin-top: 0.2rem;
        }

        @keyframes fadeInDown {
            from {
                opacity: 0;
                transform: translateY(-20px);
            }
            to {
                opacity: 1;
                transform: translateY(0);
            }
        }

        @keyframes fadeInUp {
            from {
                opacity: 0;
                transform: translateY(20px);
            }
            to {
                opacity: 1;
                transform: translateY(0);
            }
        }

        /* 卡片动画延迟 */
        .card:nth-child(1) { animation-delay: 0.1s; }
        .card:nth-child(2) { animation-delay: 0.2s; }
        .card:nth-child(3) { animation-delay: 0.3s; }
        </style>
        """,
        unsafe_allow_html=True
    )


def render_header():
    """渲染页面头部"""
    st.markdown(
        """
        <div class="app-header">
            <h1>🔒 定义词脱敏工具</h1>
            <p>智能文本脱敏 · 支持正则匹配 · 本地加密还原 · 数据不离开您的设备</p>
        </div>
        """,
        unsafe_allow_html=True
    )


def render_masking_card():
    """渲染脱敏功能卡片"""
    st.markdown('<div class="card">', unsafe_allow_html=True)
    st.markdown('<div class="card-title">📝 文本脱敏</div>', unsafe_allow_html=True)

    # 输入文本
    source_text = st.text_area(
        "原始文本",
        height=180,
        placeholder="在这里粘贴需要脱敏的内容...",
        value=Config.DEFAULT_TEXT_SAMPLE,
        label_visibility="collapsed"
    )

    # 文件上传
    source_file = st.file_uploader(
        "或上传文件（txt / docx / pdf）",
        type=["txt", "docx", "pdf"],
        label_visibility="visible"
    )

    # 高级选项
    with st.expander("⚙️ 高级选项", expanded=False):
        col1, col2 = st.columns(2)
        with col1:
            mask_mode = st.selectbox(
                "脱敏模式",
                options=[("全量替换", MaskMode.FULL), ("部分遮蔽", MaskMode.PARTIAL)],
                format_func=lambda x: x[0],
                index=0
            )[1]

            if mask_mode == MaskMode.PARTIAL:
                preserve_chars = st.slider(
                    "保留字符数",
                    min_value=1,
                    max_value=5,
                    value=1,
                    help="保留关键词开头的字符数量"
                )
            else:
                preserve_chars = 1

            mask_char = st.text_input(
                "脱敏字符",
                value="*",
                max_chars=1,
                help="用于替换敏感内容的字符"
            )

        with col2:
            enable_smart = st.checkbox(
                "🤖 启用智能识别",
                value=False,
                help="自动识别常见敏感信息（手机号、身份证、邮箱、银行卡等）"
            )

            if enable_smart:
                st.markdown("**智能识别模式：**")
                for name, pattern in PREDEFINED_PATTERNS.items():
                    st.markdown(f"- {name}：{pattern.description}")

    # 关键词输入
    keywords = st.text_area(
        "脱敏关键词（支持换行、逗号、分号分隔）",
        height=100,
        placeholder="例如：张三, 13800138000, zhangsan@example.com",
        help="输入需要脱敏的关键词，每行一个或用逗号分隔"
    )

    # 密码输入
    password = st.text_input(
        "🔑 还原密码",
        type="password",
        placeholder="设置密码用于加密原始文本（必须牢记，丢失无法找回）",
        help="密码用于加密原始文本，解密时需要使用相同密码"
    )

    # 执行按钮
    col1, col2 = st.columns([3, 1])
    with col1:
        run_button = st.button("🚀 开始脱敏", use_container_width=True, type="primary")

    st.markdown("</div>", unsafe_allow_html=True)

    return source_text, source_file, keywords, password, mask_mode, preserve_chars, mask_char, enable_smart, run_button


def render_result_card():
    """渲染结果卡片"""
    st.markdown('<div class="card">', unsafe_allow_html=True)
    st.markdown('<div class="card-title">📦 脱敏结果</div>', unsafe_allow_html=True)
    result_area = st.empty()
    st.markdown('<div class="info-box">💡 生成后可下载包含脱敏文档和加密还原文件的压缩包</div>', unsafe_allow_html=True)
    st.markdown("</div>", unsafe_allow_html=True)
    return result_area


def render_restore_card():
    """渲染还原解密卡片"""
    st.markdown('<div class="card secondary-btn">', unsafe_allow_html=True)
    st.markdown('<div class="card-title">🔓 还原原文</div>', unsafe_allow_html=True)

    restore_file = st.file_uploader(
        "选择加密还原文件（JSON）",
        type=["json"],
        key="restore_file"
    )

    restore_password = st.text_input(
        "🔑 还原密码",
        type="password",
        placeholder="输入生成时设置的密码",
        key="restore_password"
    )

    restore_button = st.button("🔓 解密并下载原文", use_container_width=True)
    restore_status = st.empty()

    st.markdown("</div>", unsafe_allow_html=True)

    return restore_file, restore_password, restore_button, restore_status


def display_stats(stats: dict):
    """显示脱敏统计信息"""
    if not stats:
        return

    st.markdown('<div class="stats-grid">', unsafe_allow_html=True)

    col_count = 0
    if stats.get("manual_keywords", 0) > 0:
        st.markdown(
            f'<div class="stat-item"><div class="stat-value">{stats["manual_keywords"]}</div><div class="stat-label">关键词</div></div>',
            unsafe_allow_html=True
        )
        col_count += 1

    for name, count in stats.get("smart_detection", {}).items():
        if count > 0:
            st.markdown(
                f'<div class="stat-item"><div class="stat-value">{count}</div><div class="stat-label">{name}</div></div>',
                unsafe_allow_html=True
            )
            col_count += 1

    if col_count == 0:
        st.markdown('<div class="info-box">未检测到需要脱敏的内容</div>', unsafe_allow_html=True)

    st.markdown('</div>', unsafe_allow_html=True)


# ============= 主应用 =============

def main():
    """主应用入口"""
    init_page_style()
    render_header()

    # 三列布局
    col1, col2, col3 = st.columns([1.3, 1, 1])

    # 左列：脱敏功能
    with col1:
        source_text, source_file, keywords, password, mask_mode, preserve_chars, mask_char, enable_smart, run_button = render_masking_card()
        mask_status = st.empty()

    # 中列：结果展示
    with col2:
        result_area = render_result_card()

    # 右列：还原功能
    with col3:
        restore_file, restore_password, restore_button, restore_status = render_restore_card()

    # ========== 处理脱敏请求 ==========
    if run_button:
        mask_status.empty()
        with mask_status:
            st.info("⏳ 处理中...")

        try:
            # 获取文本
            text = source_text.strip()
            if source_file is not None and not text:
                text = extract_file_text(source_file)
                st.success(f"✅ 已读取文件：{source_file.name}")

            if not text:
                st.error("❌ 请输入或上传需要脱敏的文本")
                st.stop()

            # 解析关键词
            keywords_list = normalize_keywords(keywords)

            if not keywords_list and not enable_smart:
                st.error("❌ 请输入关键词或启用智能识别")
                st.stop()

            if not password:
                st.error("❌ 请设置还原密码")
                st.stop()

            if len(password) < 6:
                st.error("❌ 密码长度至少6位")
                st.stop()

            # 执行脱敏
            masked, stats = build_masked_text(
                text,
                keywords_list,
                mask_mode,
                preserve_chars,
                mask_char,
                enable_smart
            )

            # 加密原文
            encrypted = encrypt_text(text, password, keywords_list)
            encrypted_dict = asdict(encrypted)

            # 生成文件
            stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            masked_docx = build_docx_bytes(masked)
            bundle_bytes = build_zip_bundle(masked_docx, encrypted_dict, stamp)

            # 显示结果
            with result_area:
                st.markdown('<div class="card">', unsafe_allow_html=True)
                st.markdown('<div class="card-title">✅ 脱敏完成</div>', unsafe_allow_html=True)

                # 显示统计
                display_stats(stats)

                # 预览脱敏结果
                with st.expander("👁️ 预览脱敏结果"):
                    st.text(masked[:500] + "..." if len(masked) > 500 else masked)

                # 下载按钮
                st.download_button(
                    label="📥 下载脱敏 + 还原包",
                    data=bundle_bytes,
                    file_name=f"masked_bundle_{stamp}.zip",
                    mime="application/zip",
                    use_container_width=True
                )

                st.markdown("</div>", unsafe_allow_html=True)

            with mask_status:
                st.success("✅ 处理完成！")

        except Exception as e:
            with mask_status:
                st.error(f"❌ 处理失败：{str(e)}")

    # ========== 处理还原请求 ==========
    if restore_button:
        with restore_status:
            st.info("⏳ 解密中...")

        try:
            if restore_file is None:
                st.error("❌ 请选择加密还原文件")
                st.stop()

            if not restore_password:
                st.error("❌ 请输入还原密码")
                st.stop()

            # 读取并解密
            payload = json.loads(restore_file.read().decode("utf-8"))
            plain = decrypt_text(payload, restore_password)

            # 生成文档
            stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            restored_docx = build_docx_bytes(plain)

            with restore_status:
                st.success("✅ 解密成功！")
                st.download_button(
                    label="📥 下载原文 DOCX",
                    data=restored_docx,
                    file_name=f"restored_{stamp}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )

        except ValueError as e:
            with restore_status:
                st.error(f"❌ 解密失败：{str(e)}")
        except Exception as e:
            with restore_status:
                st.error(f"❌ 处理失败：{str(e)}")


if __name__ == "__main__":
    main()
